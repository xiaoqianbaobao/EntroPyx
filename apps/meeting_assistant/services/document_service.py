"""
文档生成服务
Document Generation Service
"""
import logging
from typing import Dict

logger = logging.getLogger(__name__)


class DocumentService:
    """文档生成服务基类"""
    
    def generate_markdown(self, summary_data: Dict) -> str:
        """生成Markdown文档"""
        raise NotImplementedError("Subclass must implement this method")
    
    def generate_pdf(self, summary_data: Dict) -> bytes:
        """生成PDF文档"""
        raise NotImplementedError("Subclass must implement this method")
    
    def generate_docx(self, summary_data: Dict) -> bytes:
        """生成Word文档"""
        raise NotImplementedError("Subclass must implement this method")


class SimpleDocumentService(DocumentService):
    """简单的文档生成服务"""
    
    def __init__(self):
        self.name = "Simple Document Service"
    
    def generate_markdown(self, summary: object) -> str:
        """生成Markdown文档"""
        md = f"""# {summary.title}

## 基本信息

- **仓库**: {summary.repository.name}
- **会议时间**: {summary.recording.meeting_date}
- **参会人员**: {summary.recording.participants}
- **生成时间**: {summary.generated_at.strftime('%Y-%m-%d %H:%M:%S')}

## 会议摘要

{summary.summary_text}

## 讨论要点

"""
        for point in summary.key_points:
            md += f"- {point}\n"
        
        md += "\n## 决策事项\n\n"
        for decision in summary.decisions:
            md += f"- ✅ {decision}\n"
        
        if summary.action_items:
            md += "\n## 待办任务\n\n"
            md += "| 任务 | 负责人 | 截止时间 |\n"
            md += "|------|--------|----------|\n"
            for item in summary.action_items:
                task = item.get('task', '')
                assignee = item.get('assignee', '')
                deadline = item.get('deadline', '')
                md += f"| {task} | {assignee} | {deadline} |\n"
        
        md += "\n## 评审意见\n\n"
        for opinion in summary.opinions.all():
            emoji_map = {
                'issue': '🔴',
                'suggestion': '🟡',
                'decision': '🟢',
                'risk': '⚠️',
                'positive': '✨'
            }
            emoji = emoji_map.get(opinion.opinion_type, '📝')
            status = '✓ 已解决' if opinion.is_resolved else '○ 待解决'
            md += f"{emoji} **{opinion.get_opinion_type_display()}**: {opinion.content}\n"
            md += f"  - 优先级: {opinion.get_priority_display()}\n"
            md += f"  - 状态: {status}\n\n"
        
        return md
    
    def generate_pdf(self, summary: object) -> bytes:
        """生成PDF文档"""
        try:
            # 先生成Markdown
            md_content = self.generate_markdown(summary)
            
            # 转换为HTML
            from markdown2 import markdown
            html_content = f"""
            <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: "Microsoft YaHei", Arial, sans-serif; padding: 20px; }}
                    h1 {{ color: #333; border-bottom: 2px solid #0066cc; }}
                    h2 {{ color: #0066cc; }}
                    table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                </style>
            </head>
            <body>
                {markdown(md_content, extras=['tables'])}
            </body>
            </html>
            """
            
            # 转换为PDF
            from weasyprint import HTML
            pdf_bytes = HTML(string=html_content).write_pdf()
            
            return pdf_bytes
            
        except ImportError:
            logger.warning("weasyprint not installed, using markdown content")
            # 返回Markdown内容
            return self.generate_markdown(summary).encode('utf-8')
        except Exception as e:
            logger.error(f"PDF generation failed: {str(e)}")
            # 降级到Markdown
            return self.generate_markdown(summary).encode('utf-8')
    
    def generate_docx(self, summary: object) -> bytes:
        """生成Word文档"""
        try:
            from docx import Document
            from docx.shared import Pt
            from io import BytesIO
            
            doc = Document()
            
            # 标题
            title = doc.add_heading(summary.title, level=1)
            
            # 基本信息
            doc.add_heading('基本信息', level=2)
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Light Grid Accent 1'
            
            cells = table.rows[0].cells
            cells[0].text = '仓库'
            cells[1].text = summary.repository.name
            
            cells = table.rows[1].cells
            cells[0].text = '会议时间'
            cells[1].text = str(summary.recording.meeting_date)
            
            cells = table.rows[2].cells
            cells[0].text = '参会人员'
            cells[1].text = summary.recording.participants
            
            cells = table.rows[3].cells
            cells[0].text = '生成时间'
            cells[1].text = summary.generated_at.strftime('%Y-%m-%d %H:%M:%S')
            
            # 会议摘要
            doc.add_heading('会议摘要', level=2)
            doc.add_paragraph(summary.summary_text)
            
            # 讨论要点
            doc.add_heading('讨论要点', level=2)
            for point in summary.key_points:
                doc.add_paragraph(point, style='List Bullet')
            
            # 决策事项
            doc.add_heading('决策事项', level=2)
            for decision in summary.decisions:
                doc.add_paragraph(decision, style='List Bullet')
            
            # 待办任务
            if summary.action_items:
                doc.add_heading('待办任务', level=2)
                task_table = doc.add_table(rows=len(summary.action_items)+1, cols=3)
                task_table.style = 'Light Grid Accent 1'
                
                hdr_cells = task_table.rows[0].cells
                hdr_cells[0].text = '任务'
                hdr_cells[1].text = '负责人'
                hdr_cells[2].text = '截止时间'
                
                for idx, item in enumerate(summary.action_items):
                    row_cells = task_table.rows[idx+1].cells
                    row_cells[0].text = item.get('task', '')
                    row_cells[1].text = item.get('assignee', '')
                    row_cells[2].text = item.get('deadline', '')
            
            # 评审意见
            doc.add_heading('评审意见', level=2)
            for opinion in summary.opinions.all():
                p = doc.add_paragraph()
                p.add_run(f"{opinion.get_opinion_type_display()}: ").bold = True
                p.add_run(opinion.content)
                p = doc.add_paragraph(f"  优先级: {opinion.get_priority_display()}, 状态: {'已解决' if opinion.is_resolved else '待解决'}", style='List Bullet')
            
            # 保存到内存
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            return buffer.read()
            
        except ImportError:
            logger.warning("python-docx not installed, using markdown content")
            return self.generate_markdown(summary).encode('utf-8')
        except Exception as e:
            logger.error(f"DOCX generation failed: {str(e)}")
            return self.generate_markdown(summary).encode('utf-8')


def get_document_service(service_type='simple'):
    """
    获取文档生成服务实例
    
    Args:
        service_type: 服务类型 ('simple')
    
    Returns:
        DocumentService实例
    """
    return SimpleDocumentService()