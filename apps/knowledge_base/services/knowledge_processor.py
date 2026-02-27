"""
知识库文档处理服务
支持PDF、MD、Word等格式的文档解析和知识提取
"""
import os
import logging
import PyPDF2
import docx
from bs4 import BeautifulSoup
import re
from django.conf import settings

logger = logging.getLogger(__name__)

class KnowledgeProcessor:
    """知识库文档处理器"""
    
    def __init__(self):
        self.supported_formats = {
            'pdf': self._process_pdf,
            'md': self._process_markdown,
            'markdown': self._process_markdown,
            'txt': self._process_text,
            'docx': self._process_word,
            'doc': self._process_word
        }
    
    def process_document(self, file_path, file_name, file_type):
        """
        处理文档并提取知识
        
        Args:
            file_path: 文件路径
            file_name: 文件名
            file_type: 文件类型
            
        Returns:
            dict: 处理结果
        """
        try:
            # 获取处理器
            processor = self.supported_formats.get(file_type.lower())
            if not processor:
                raise ValueError(f"不支持的文件格式: {file_type}")
            
            # 处理文档
            content = processor(file_path)
            
            # 提取结构化信息
            structured_data = self._extract_structured_data(content, file_name)
            
            return {
                'success': True,
                'content': content,
                'structured_data': structured_data,
                'file_name': file_name,
                'file_type': file_type,
                'content_length': len(content)
            }
            
        except Exception as e:
            logger.error(f"处理文档失败 {file_name}: {str(e)}")
            return {
                'success': False,
                'error': str(e),
                'file_name': file_name
            }
    
    def _process_pdf(self, file_path):
        """处理PDF文件"""
        content = []
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        content.append(f"--- 第{page_num + 1}页 ---\n{text}")
                
                logger.info(f"PDF处理完成，共{len(pdf_reader.pages)}页")
                
        except Exception as e:
            logger.error(f"PDF处理失败: {str(e)}")
            raise
        
        return '\n\n'.join(content)
    
    def _process_markdown(self, file_path):
        """处理Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            # 清理Markdown格式，保留纯文本
            soup = BeautifulSoup(content, "html.parser")
            text = soup.get_text()
            
            # 移除多余的空白行
            text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
            
            logger.info("Markdown处理完成")
            return text
            
        except Exception as e:
            logger.error(f"Markdown处理失败: {str(e)}")
            raise
    
    def _process_text(self, file_path):
        """处理文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()
            
            logger.info("文本文件处理完成")
            return content
            
        except Exception as e:
            logger.error(f"文本文件处理失败: {str(e)}")
            raise
    
    def _process_word(self, file_path):
        """处理Word文档"""
        try:
            doc = docx.Document(file_path)
            content = []
            
            # 处理段落
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text)
            
            # 处理表格
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    
                    if row_text:
                        content.append(" | ".join(row_text))
            
            logger.info("Word文档处理完成")
            return '\n\n'.join(content)
            
        except Exception as e:
            logger.error(f"Word文档处理失败: {str(e)}")
            raise
    
    def _extract_structured_data(self, content, file_name):
        """提取结构化数据"""
        structured_data = {
            'title': file_name,
            'sections': [],
            'keywords': [],
            'entities': [],
            'summary': ''
        }
        
        try:
            # 提取标题（如果有）
            lines = content.split('\n')
            for i, line in enumerate(lines[:10]):  # 检查前10行
                line = line.strip()
                if line and not line.startswith('-') and not line.startswith('*'):
                    structured_data['title'] = line
                    break
            
            # 提取章节
            section_pattern = r'^(#{1,6})\s+(.+)$'
            for i, line in enumerate(lines):
                match = re.match(section_pattern, line)
                if match:
                    level = len(match.group(1))
                    title = match.group(2).strip()
                    structured_data['sections'].append({
                        'level': level,
                        'title': title,
                        'index': i
                    })
            
            # 提取关键词（简化版）
            words = re.findall(r'\b[a-zA-Z\u4e00-\u9fff]{2,}\b', content)
            word_freq = {}
            for word in words:
                if len(word) > 2:
                    word_freq[word] = word_freq.get(word, 0) + 1
            
            # 按频率排序
            sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
            structured_data['keywords'] = [word for word, freq in sorted_words[:20]]
            
            # 提取实体（简化版）
            entity_pattern = r'([A-Z][a-z]+(?:\s+[A-Z][a-z]+)*)|([一二三四五六七八九十百千万]+)'
            entities = re.findall(entity_pattern, content)
            structured_data['entities'] = list(set([e[0] or e[1] for e in entities if e[0] or e[1]]))[:10]
            
            # 生成摘要
            sentences = re.split(r'[。！？\n]', content)
            sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
            if sentences:
                structured_data['summary'] = sentences[0][:200] + '...' if len(sentences[0]) > 200 else sentences[0]
            
        except Exception as e:
            logger.error(f"提取结构化数据失败: {str(e)}")
        
        return structured_data
    
    def validate_document(self, file_path, file_type):
        """验证文档格式"""
        if not os.path.exists(file_path):
            raise FileNotFoundError("文件不存在")
        
        if not file_type.lower() in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_type}")
        
        # 检查文件大小（限制为100MB）
        file_size = os.path.getsize(file_path)
        if file_size > 100 * 1024 * 1024:
            raise ValueError("文件大小超过100MB限制")
        
        return True